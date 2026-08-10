from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\gaura\OneDrive\Desktop\AI-NLP-COPO-MAPPER")
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "AI_COPO_Major_Project_Logbook.docx"
LOGO_PATH = ROOT / "_tmp_logbook_extracted" / "word" / "media" / "image1.jpeg"


PROJECT_TITLE = "AI-Based Automated CO-PO Mapping System using NLP and Machine Learning"
ACADEMIC_YEAR = "2025-2026"
COURSE_CODE = "CSP701 / CSP801"
GUIDE_NAME = "Prof. Avinash Gondal"
CO_GUIDE_NAME = ""
DEPARTMENT = "Department of Computer Engineering"
INSTITUTE = "WATUMULL INSTITUTE OF ENGINEERING & TECHNOLOGY"
SUBJECT_A = "Project Based Learning - Major Project - A"
SUBJECT_B = "Project Based Learning - Major Project - B"

TEAM_MEMBERS = [
    ("1", "Gaurav Anil Zambare", "________"),
    ("2", "Priyanka Pravin Solse", "________"),
    ("3", "Hemanshu Sushilkumar Raut", "________"),
    ("4", "Pankaj Kamlesh Gehlot", "________"),
]

MISSION = "To provide high quality of technical education so that our engineers develop innovative research capabilities and become leaders and decision makers."
VISION_DEPT = "To become a center of excellence in technical education and to produce globally competent professionals who excel in emerging areas of computer engineering by imparting relevant knowledge and skills and inculcating ethical human values."
MISSION_DEPT = [
    "M1: Be in the forefront of the domain of computer engineering through national and international recognition in research and education.",
    "M2: Foster development of problem solving and communication skills as an integral part of the profession to solve real-world problems, thus amplifying the potential for lifelong learning.",
    "M3: Impart quality learning experience through effective classroom practices and meaningful interaction between student and teacher.",
]

PO_TABLE = [
    ("PO1 Engineering Knowledge", "Apply the knowledge of mathematics, science, engineering fundamentals, and engineering specialization to the solution of complex engineering problems."),
    ("PO2 Problem Analysis", "Identify, formulate, review research literature, and analyze complex engineering problems reaching substantiated conclusions using first principles of mathematics, natural sciences, and engineering sciences."),
    ("PO3 Design / Development of Solutions", "Design solutions for complex engineering problems and design system components or processes that meet specified needs with appropriate consideration for public health, safety, and societal factors."),
    ("PO4 Conduct Investigations of Complex Problems", "Use research-based knowledge and research methods including design of experiments, analysis and interpretation of data, and synthesis of information to provide valid conclusions."),
    ("PO5 Modern Tool Usage", "Create, select, and apply appropriate techniques, resources, and modern engineering and IT tools, including prediction and modeling, to complex engineering activities with an understanding of limitations."),
    ("PO6 The Engineer and Society", "Apply reasoning informed by contextual knowledge to assess societal, health, safety, legal, and cultural issues and responsibilities relevant to professional engineering practice."),
    ("PO7 Environment and Sustainability", "Understand the impact of professional engineering solutions in societal and environmental contexts and demonstrate knowledge of sustainable development."),
    ("PO8 Ethics", "Apply ethical principles and commit to professional ethics, responsibilities, and norms of engineering practice."),
    ("PO9 Individual and Team Work", "Function effectively as an individual, and as a member or leader in diverse teams and multidisciplinary settings."),
    ("PO10 Communication", "Communicate effectively on complex engineering activities with the engineering community and society through reports, documentation, presentations, and clear instructions."),
    ("PO11 Project Management and Finance", "Demonstrate knowledge and understanding of engineering and management principles and apply these to one’s own work as a member and leader to manage projects in multidisciplinary environments."),
]

PSO_TABLE = [
    ("PSO I", "Develop coding skills and work in multidisciplinary environments to solve computing problems using suitable hardware and software tools."),
    ("PSO II", "Understand and apply knowledge in diverse areas of computer engineering such as algorithms, networking, databases, web technologies, artificial intelligence, and data-driven systems."),
    ("PSO III", "Apply standard industry practices, tools, and technologies in the development of practical software systems and intelligent engineering solutions."),
]

COURSE_OBJECTIVE = (
    "The project work facilitates students to apply technical, professional, and ethical skills developed during graduation for solving a real academic problem through requirement analysis, design, implementation, testing, and validation."
)

COURSE_OUTCOMES = [
    "Develop understanding of Outcome Based Education, CO-PO mapping methodology, and the academic problem domain through detailed literature survey and requirement gathering.",
    "Identify, analyze, and structure Course Outcomes, Program Outcomes, Working Knowledge statements, and faculty validation requirements for an automated mapping system.",
    "Design and implement an AI-assisted mapping engine using TF-IDF, BERT-based semantic similarity, Bloom’s Taxonomy analysis, and explainable justification generation.",
    "Develop a complete software solution with modules for dashboard mapping, analytics, question-to-CO detection, attainment calculation, PDF report generation, and data persistence.",
    "Evaluate mapping quality through faculty feedback, teacher reference documents, and machine learning enhancement using XGBoost-based reranking.",
    "Demonstrate teamwork, documentation, presentation, and project management skills through successful completion, testing, and deployment of the proposed system.",
]

PROJECT_ABSTRACT = (
    "The proposed project automates Course Outcome to Program Outcome mapping using Natural Language Processing and Machine Learning. "
    "Traditional CO-PO mapping in academic institutions is manual, time-consuming, subjective, and inconsistent across faculty members. "
    "The developed system accepts Course Outcomes as input and compares them with AICTE Program Outcomes and Working Knowledge statements using text preprocessing, TF-IDF similarity, lexical overlap, BERT semantic similarity, and Bloom’s Taxonomy detection. "
    "The system further improves mapping quality using teacher reference documents and an XGBoost-based hybrid reranking model trained on validated subject files. "
    "Additional modules include question-to-CO detection, printable question paper generation, attainment calculation for both C Scheme and NEP Scheme, analytics dashboards, explainable justification generation, and PDF reporting. "
    "The final solution reduces manual effort, increases standardization, and provides faculty-friendly editable outputs for academic documentation."
)

SEM7_TASKS = [
    "Identified the academic problem of manual and inconsistent CO-PO mapping in OBE-based university documentation.",
    "Conducted literature survey on NLP-based text similarity, outcome-based education, Bloom’s Taxonomy, and explainable academic mapping systems.",
    "Finalized project title, system scope, and high-level module architecture for dashboard, analytics, attainment, and question detection.",
    "Collected AICTE PO statements, Working Knowledge statements, C Scheme and NEP subject CO data, and initial faculty requirements.",
    "Designed database schema using MongoDB for users, subjects, mappings, feedback, teacher references, and generated reports.",
    "Implemented text preprocessing, token normalization, lexical overlap baseline, and TF-IDF similarity prototype.",
    "Integrated BERT-based semantic similarity and hybrid score calculation for CO-PO mapping.",
    "Implemented Bloom’s Taxonomy keyword mapping and AI level / confidence generation logic.",
    "Developed dashboard and result pages for mapping generation, editable justification, and PDF report output.",
    "Implemented faculty validation, reason selection, matrix editing, and storage of mapping decisions in MongoDB.",
    "Prepared semester VII synopsis, module breakdown, draft results, and review presentation.",
    "Completed semester VII documentation and internal review corrections for continuation into final implementation phase.",
]

SEM8_TASKS = [
    "Implemented question-to-CO detection flow and printable question paper generation with BTL and CO mapping.",
    "Added secure PDF generation, institutional branding, watermark handling, and same-page printable previews.",
    "Developed attainment module with Excel parsing, CO-wise question analysis, and scheme-based level calculation for C Scheme and NEP Scheme.",
    "Built analytics page for CO-PO visualization, Bloom distribution, and subject-wise outcome summaries.",
    "Collected teacher reference files from multiple subjects and created old 12-PO to new 11-PO conversion mapping.",
    "Built training dataset generation pipeline from teacher CO-PO mapping and justification documents.",
    "Trained XGBoost model on teacher reference dataset and evaluated mapping accuracy and F1-score.",
    "Integrated trained model into the live NLP pipeline using hybrid reranking with TF-IDF, BERT, lexical, and Bloom features.",
    "Enhanced OCR/PDF extraction utilities, subject-catalog loading, spell correction, and workflow stabilization.",
    "Resolved UI issues related to persistence, printable reports, sorting, subject selection, and editable justification handling.",
    "Prepared black book, IEEE implementation paper, weekly progress records, and demonstration materials.",
    "Completed final integration testing, faculty walkthrough, viva preparation, and project submission readiness.",
]

WEEKLY_PLAN = [
    ("1", "Project domain discussion, problem identification, and understanding of manual CO-PO mapping issues."),
    ("2", "Finalization of project title, objective draft, and study of Outcome Based Education concepts."),
    ("3", "Literature survey on NLP, semantic similarity, Bloom’s Taxonomy, and academic automation systems."),
    ("4", "Collection of AICTE PO statements, WK statements, and sample university CO documents."),
    ("5", "Requirement analysis with faculty expectations, mapping workflow study, and module planning."),
    ("6", "System architecture design, frontend-backend planning, and database schema preparation."),
    ("7", "Collection and cleaning of C Scheme and NEP Scheme subject data for MongoDB storage."),
    ("8", "Initial dashboard design and development of course outcome input, mapping mode, and result workflow."),
    ("9", "Implementation of text preprocessing, token cleaning, and baseline lexical / TF-IDF similarity."),
    ("10", "Integration of BERT semantic similarity and hybrid scoring strategy for CO-PO matching."),
    ("11", "Bloom’s Taxonomy detection, AI level, confidence labeling, and justification logic design."),
    ("12", "Development of result page, editable justification handling, and matrix-based faculty validation flow."),
    ("13", "Review of semester VII progress, synopsis preparation, and consolidation of first phase work."),
    ("14", "Detailed planning for implementation phase, bug review, and finalization of implementation roadmap."),
    ("15", "Start of implementation phase: live backend integration of CO-PO mapping pipeline."),
    ("16", "Implementation of analytics page, subject-wise mapping visualization, and bloom distribution charts."),
    ("17", "Implementation of question-to-CO detection and question paper structure builder."),
    ("18", "Question paper printable PDF generation, college branding, and watermark handling."),
    ("19", "Implementation of attainment module with Excel parsing and question-wise CO analysis."),
    ("20", "Addition of C Scheme and NEP Scheme attainment calculation rules and formatted report generation."),
    ("21", "Implementation of dashboard extraction, subject catalog loading, and workflow stability fixes."),
    ("22", "Collection of teacher reference files and preparation of old 12-PO to new 11-PO conversion logic."),
    ("23", "Creation of training dataset and XGBoost-based model training for teacher-informed hybrid mapping."),
    ("24", "Integration of trained model into live hybrid system, tuning, and validation against reference data."),
    ("25", "Final testing, PDF/report polishing, log book preparation, and black book / IEEE paper drafting."),
    ("26", "Project demonstration readiness, final review corrections, viva preparation, and submission documentation."),
]

PROJECT_PLANNER = [
    ("1", "Selection of Topic", "July"),
    ("2", "Literature Survey", "July-Aug"),
    ("3", "Requirement Gathering and Dataset Collection", "Aug-Sep"),
    ("4", "Implementation - I (NLP Mapping Engine)", "Sep-Oct"),
    ("5", "Synopsis / Semester VII Review", "Nov"),
    ("6", "Implementation - II (Frontend, Analytics, Question Detection)", "Dec-Jan"),
    ("7", "Teacher Reference Data Processing and Model Training", "Jan-Feb"),
    ("8", "Attainment, PDF, and Security Enhancements", "Feb-Mar"),
    ("9", "Final Report, Black Book, and IEEE Paper", "Mar-Apr"),
    ("10", "Final Review, Demo, and Submission", "Apr"),
]

RUBRIC_ROWS = [
    ("Task Completion", "Completion of assigned work and submission to guide"),
    ("Preparedness", "Understanding of the current project phase and technical readiness"),
    ("Team Work", "Participation, collaboration, and contribution within the group"),
    ("Punctuality", "Regularity, discipline, and on-time completion of weekly tasks"),
]


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def add_title(document, text, size=18, bold=True, spacing_after=6):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)


def add_logo(document):
    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(6.0))


def add_section_heading(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)


def add_normal_paragraph(document, text, bold_prefix=None):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(11)
        run2 = p.add_run(text[len(bold_prefix):])
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return p


def add_members_table(document):
    add_section_heading(document, "Project Team Details")
    table = document.add_table(rows=1, cols=3)
    format_table(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Sr. No.", True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[1], "Name of Student", True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[2], "Roll No. (Editable)", True, WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr:
        shade_cell(c, "D9EAF7")

    for sr, name, roll in TEAM_MEMBERS:
        row = table.add_row().cells
        set_cell_text(row[0], sr, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], name)
        set_cell_text(row[2], roll, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_two_col_table(document, title, rows):
    add_section_heading(document, title)
    table = document.add_table(rows=0, cols=2)
    format_table(table)
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key, True)
        set_cell_text(row[1], value)
    return table


def add_course_outcomes_table(document):
    add_section_heading(document, "Course Objectives and Course Outcomes")
    add_normal_paragraph(document, f"Course Objective: {COURSE_OBJECTIVE}", bold_prefix="Course Objective:")
    table = document.add_table(rows=1, cols=2)
    format_table(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "CO No.", True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[1], "Course Outcome", True, WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr:
        shade_cell(c, "D9EAF7")
    for index, co in enumerate(COURSE_OUTCOMES, start=1):
        row = table.add_row().cells
        set_cell_text(row[0], f"CO{index}", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], co)


def add_project_planner(document):
    add_section_heading(document, "Project Planner")
    table = document.add_table(rows=1, cols=3)
    format_table(table)
    headers = ["Sr. No.", "Assignment / Milestone", "Planned Timeline"]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for sr, assignment, timeline in PROJECT_PLANNER:
        row = table.add_row().cells
        set_cell_text(row[0], sr, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], assignment)
        set_cell_text(row[2], timeline, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_weekwise_plan_table(document):
    add_section_heading(document, "Week-wise Work Distribution (Week 1 to Week 26)")
    table = document.add_table(rows=1, cols=2)
    format_table(table)
    set_cell_text(table.rows[0].cells[0], "Week No.", True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[0].cells[1], "Planned Work / Progress Activity", True, WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(table.rows[0].cells[0], "D9EAF7")
    shade_cell(table.rows[0].cells[1], "D9EAF7")

    for week_no, task in WEEKLY_PLAN:
        row = table.add_row().cells
        set_cell_text(row[0], week_no, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], task)


def add_task_summary(document, title, tasks):
    add_section_heading(document, title)
    for idx, task in enumerate(tasks, start=1):
        p = document.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{idx}) {task}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_weekly_assessment_table(document, heading, member_name, start_week, end_week):
    add_section_heading(document, heading)
    add_normal_paragraph(document, f"Student Name: {member_name}", bold_prefix="Student Name:")
    add_normal_paragraph(document, "Roll No.: ____________________", bold_prefix="Roll No.:")
    table = document.add_table(rows=1, cols=8)
    format_table(table)
    headers = [
        "Week No.",
        "Date",
        "Task Completion (5)",
        "Preparedness (5)",
        "Team Work (5)",
        "Punctuality (5)",
        "Total (20)",
        "Signature of Guide",
    ]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, True, WD_ALIGN_PARAGRAPH.CENTER, 9)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")

    for week in range(start_week, end_week + 1):
        row = table.add_row().cells
        set_cell_text(row[0], str(week), align=WD_ALIGN_PARAGRAPH.CENTER)
        for col in range(1, 8):
            set_cell_text(row[col], "", align=WD_ALIGN_PARAGRAPH.CENTER)


def add_attendance_table(document, heading, rows_count):
    add_section_heading(document, heading)
    table = document.add_table(rows=1, cols=5)
    format_table(table)
    headers = ["Date"] + [f"{name} - Sign / Absent" for _, name, _ in TEAM_MEMBERS]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, True, WD_ALIGN_PARAGRAPH.CENTER, 9)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for _ in range(rows_count):
        row = table.add_row().cells
        for idx in range(5):
            set_cell_text(row[idx], "", align=WD_ALIGN_PARAGRAPH.CENTER)


def add_rubric_table(document):
    add_section_heading(document, "Assessment Rubrics")
    table = document.add_table(rows=1, cols=2)
    format_table(table)
    set_cell_text(table.rows[0].cells[0], "Assessment Outcome", True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[0].cells[1], "Description", True, WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(table.rows[0].cells[0], "D9EAF7")
    shade_cell(table.rows[0].cells[1], "D9EAF7")
    for name, desc in RUBRIC_ROWS:
        row = table.add_row().cells
        set_cell_text(row[0], name, True)
        set_cell_text(row[1], desc)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    add_logo(doc)
    add_title(doc, DEPARTMENT, 16)
    add_title(doc, "PROJECT LOG BOOK", 18)
    add_normal_paragraph(doc, f"Course code: {COURSE_CODE}", bold_prefix="Course code:")
    add_normal_paragraph(doc, f"Academic Year: {ACADEMIC_YEAR}", bold_prefix="Academic Year:")
    add_normal_paragraph(doc, f"Title of the Project: {PROJECT_TITLE}", bold_prefix="Title of the Project:")
    add_normal_paragraph(doc, f"Name of Guide: {GUIDE_NAME}", bold_prefix="Name of Guide:")
    add_normal_paragraph(doc, f"Name of Co-Guide (If Any): {CO_GUIDE_NAME}", bold_prefix="Name of Co-Guide (If Any):")
    add_members_table(doc)

    doc.add_page_break()
    add_section_heading(doc, "Mission of the Institute")
    add_normal_paragraph(doc, MISSION)
    add_section_heading(doc, "Vision of Computer Engineering Department")
    add_normal_paragraph(doc, VISION_DEPT)
    add_section_heading(doc, "Mission of Computer Engineering Department")
    for line in MISSION_DEPT:
        add_normal_paragraph(doc, line)
    add_two_col_table(doc, "Program Outcomes (PO)", PO_TABLE)
    add_two_col_table(doc, "Program Specific Outcomes (PSO)", PSO_TABLE)

    doc.add_page_break()
    add_normal_paragraph(doc, f"B.E. Computer Engineering - Course code: {COURSE_CODE}", bold_prefix="B.E. Computer Engineering - Course code:")
    add_normal_paragraph(doc, f"Subject: {SUBJECT_A}", bold_prefix="Subject:")
    add_normal_paragraph(doc, f"Subject: {SUBJECT_B}", bold_prefix="Subject:")
    add_course_outcomes_table(doc)
    add_section_heading(doc, "Project Abstract")
    add_normal_paragraph(doc, PROJECT_ABSTRACT)
    add_project_planner(doc)
    add_weekwise_plan_table(doc)

    doc.add_page_break()
    add_task_summary(doc, "Semester VII Work Summary", SEM7_TASKS)
    add_task_summary(doc, "Semester VIII Work Summary", SEM8_TASKS)
    add_rubric_table(doc)

    for _, member_name, _ in TEAM_MEMBERS:
        doc.add_page_break()
        add_weekly_assessment_table(doc, "Major Project - A Weekly Assessment Sheet (Week 1 to Week 14)", member_name, 1, 14)

    doc.add_page_break()
    add_attendance_table(doc, "Attendance Record - Major Project A", 14)

    for _, member_name, _ in TEAM_MEMBERS:
        doc.add_page_break()
        add_weekly_assessment_table(doc, "Major Project - B Weekly Assessment Sheet (Week 15 to Week 26)", member_name, 15, 26)

    doc.add_page_break()
    add_attendance_table(doc, "Attendance Record - Major Project B", 12)

    doc.add_page_break()
    add_section_heading(doc, "Final Notes")
    add_normal_paragraph(doc, "This log book is prepared for the project titled 'AI-Based Automated CO-PO Mapping System using NLP and Machine Learning'. All major project activities, reviews, testing, and deliverables may be updated by the team and guide during the academic year.")
    add_normal_paragraph(doc, f"Guide: {GUIDE_NAME}", bold_prefix="Guide:")
    add_normal_paragraph(doc, "Student Roll Numbers, dates, marks, attendance, and signatures are intentionally left editable for departmental use.")

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_document()
